import streamlit as st
from docx import Document
import io

# Configuración de la página
st.set_page_config(page_title="Sesión 2: El judaísmo", page_icon="📜", layout="centered")

# Título Principal basado en la sesión [cite: 3]
st.title("Ficha Diagnóstica: El judaísmo")
st.subheader("Cuna de las religiones monoteístas")

# Datos del estudiante
st.markdown("---")
col1, col2 = st.columns(2)
with col1:
    nombre = st.text_input("Nombres y Apellidos:")
with col2:
    grado = st.text_input("Grado y Sección (Ej. 1° A):")

st.markdown("---")

# Actividad 1 [cite: 7]
st.header("1. La raíz del monoteísmo")
st.write("Explica brevemente: ¿Por qué Abraham es una figura clave en el judaísmo y qué significa la 'alianza'?")
actividad1 = st.text_area("Escribe tu respuesta aquí...", key="act1")

# Actividad 2 [cite: 7]
st.header("2. Diálogo Respetuoso en Redes Sociales")
st.info("**Caso:** Estás navegando en tus redes sociales y ves que un usuario hace un comentario ofensivo burlándose de las creencias y prácticas de una religión distinta a la suya.")
st.write("Frente a esta situación, redacta una respuesta respetuosa para promover el diálogo y propón **una regla de convivencia** para entornos virtuales.")
actividad2 = st.text_area("Tu respuesta respetuosa y regla de convivencia...", key="act2")

# Actividad 3 [cite: 7]
st.header("3. Cuadro Comparativo")
st.write("Con base en lo conversado en clase, escribe al menos **una semejanza y una diferencia** entre el judaísmo y el cristianismo, recordando hacerlo sin prejuicios.")
actividad3 = st.text_area("Semejanzas y Diferencias...", key="act3")

st.markdown("---")

# Lógica para generar y descargar el documento Word
if st.button("Generar mi Evidencia en Word"):
    if nombre.strip() == "" or grado.strip() == "":
        st.error("⚠️ Por favor, completa tus nombres y tu grado antes de descargar.")
    else:
        # Crear el documento
        doc = Document()
        doc.add_heading('Ficha de Trabajo: El judaísmo', 0)
        
        doc.add_paragraph(f'Estudiante: {nombre}')
        doc.add_paragraph(f'Grado y Sección: {grado}')
        doc.add_paragraph('---')
        
        doc.add_heading('1. La raíz del monoteísmo', level=1)
        doc.add_paragraph(actividad1)
        
        doc.add_heading('2. Diálogo Respetuoso en Redes Sociales', level=1)
        doc.add_paragraph(actividad2)
        
        doc.add_heading('3. Cuadro Comparativo', level=1)
        doc.add_paragraph(actividad3)
        
        # Guardar el documento en memoria
        bio = io.BytesIO()
        doc.save(bio)
        
        # Botón de descarga
        st.success("¡Tu archivo está listo!")
        st.download_button(
            label="📥 Descargar archivo .docx",
            data=bio.getvalue(),
            file_name=f"Evidencia_{nombre.replace(' ', '_')}_{grado.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )